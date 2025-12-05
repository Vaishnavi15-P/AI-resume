SKILL_DB = [
    "python", "java", "javascript", "react", "node.js", 
    "fastapi", "django", "sql", "mongodb", "git", "aws"
]

def extract_skills(text):
    text = text.lower()
    skills_found = [skill for skill in SKILL_DB if skill in text]
    return skills_found

def extract_sections(text):
    text = text.lower()
    sections = {
        "skills": [],
        "experience": "",
        "education": "",
        "projects": ""
    }

    # Extract skills
    sections["skills"] = extract_skills(text)

    # Extract experience
    if "experience" in text:
        exp = text.split("experience")[-1]
        sections["experience"] = exp.split("education")[0] if "education" in exp else exp

    # Extract education
    if "education" in text:
        edu = text.split("education")[-1]
        sections["education"] = edu.split("projects")[0] if "projects" in edu else edu

    # Extract projects
    if "projects" in text:
        proj = text.split("projects")[-1]
        sections["projects"] = proj

    return sections


from fastapi import FastAPI, UploadFile, File
import shutil, os
from pdfminer.high_level import extract_text as extract_pdf_text
import docx

app = FastAPI()

UPLOAD_FOLDER = "uploads"
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

@app.get("/ping")
def ping():
    return {"message": "Backend is working!"}

@app.post("/upload/")
async def upload_resume(file: UploadFile = File(...)):
    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    return {"filename": file.filename, "message": "File uploaded successfully!"}

@app.post("/extract_text/")
async def extract_text(file: UploadFile = File(...)):
    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    text = ""
    if file.filename.endswith(".pdf"):
        text = extract_pdf_text(file_path)
    elif file.filename.endswith(".docx"):
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        text = "Unsupported file type!"
    
    return {"filename": file.filename, "text": text}

@app.post("/parse_resume/")
async def parse_resume(file: UploadFile = File(...)):
    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    text = ""
    if file.filename.endswith(".pdf"):
        text = extract_pdf_text(file_path)
    elif file.filename.endswith(".docx"):
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        return {"error": "Unsupported file type!"}

    sections = extract_sections(text)
    return {"filename": file.filename, "sections": sections}

